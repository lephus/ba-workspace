"""
Template filler: open a .docx template, locate tables by section header,
clear sample data rows, and fill with structured JSON data from the LLM.

This module does NO LLM calls — it operates purely on python-docx objects
and the validated JSON dict.
"""
import io
import uuid
from datetime import datetime

from docx import Document
from docx.table import Table

from app.services.template_registry import get_schema, get_template_path


def _header_text(table: Table) -> str:
    """Extract the first row's merged text (section header) from a table."""
    if not table.rows:
        return ""
    return table.rows[0].cells[0].text.strip().upper()


def _find_table(doc: Document, header_match: str, skip_first: int = 1) -> Table | None:
    """
    Find the table whose first row contains header_match (case-insensitive).
    Skips the first `skip_first` tables (typically the title table) to avoid
    false matches like "BUSINESS REQUIREMENTS DOCUMENT" matching before
    the actual "BUSINESS REQUIREMENTS" section.
    """
    needle = header_match.upper()
    for i, table in enumerate(doc.tables):
        if i < skip_first:
            continue
        if needle in _header_text(table):
            return table
    return None


def _clear_data_rows(table: Table, keep_rows: int = 2):
    """
    Remove all rows after the first `keep_rows` rows.
    Typically row 0 = section header, row 1 = column headers.
    """
    while len(table.rows) > keep_rows:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)


def _copy_row_style(source_row, new_row):
    """Copy cell formatting from source row to new row (font, alignment, etc.)."""
    for i, cell in enumerate(new_row.cells):
        if i < len(source_row.cells):
            src_cell = source_row.cells[i]
            if src_cell.paragraphs and cell.paragraphs:
                src_para = src_cell.paragraphs[0]
                dst_para = cell.paragraphs[0]
                if src_para.paragraph_format.alignment is not None:
                    dst_para.paragraph_format.alignment = src_para.paragraph_format.alignment
                if src_para.runs:
                    src_run = src_para.runs[0]
                    if dst_para.runs:
                        dst_run = dst_para.runs[0]
                    else:
                        dst_run = dst_para.add_run()
                    if src_run.font.size:
                        dst_run.font.size = src_run.font.size
                    if src_run.font.name:
                        dst_run.font.name = src_run.font.name


def _fill_columnar(table: Table, columns: list[str], rows_data: list[dict]):
    """
    Fill a columnar table.
    Assumes row 0 = section header (merged), row 1 = column headers.
    Clears rows 2+ and inserts new rows from rows_data.
    Skips rows that look like phase separators (all columns have same value).
    """
    _clear_data_rows(table, keep_rows=2)

    style_source = table.rows[1] if len(table.rows) > 1 else None

    for row_dict in rows_data:
        new_row = table.add_row()
        for col_idx, col_name in enumerate(columns):
            if col_idx < len(new_row.cells):
                value = str(row_dict.get(col_name, "[TBD]"))
                new_row.cells[col_idx].text = value
        if style_source:
            _copy_row_style(style_source, new_row)


def _fill_key_value(table: Table, fields: list[str], data: dict):
    """
    Fill a key_value table where each row (after header) is label | value.
    Clears existing value cells and writes new values.
    If there are more fields than existing rows, appends new rows.
    """
    existing_label_rows = {}
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        if len(row.cells) >= 2:
            label = row.cells[0].text.strip()
            existing_label_rows[label] = row_idx

    for field in fields:
        value = str(data.get(field, "[TBD]"))

        matched_idx = None
        for label, idx in existing_label_rows.items():
            if field.lower() in label.lower() or label.lower() in field.lower():
                matched_idx = idx
                break

        if matched_idx is not None:
            row = table.rows[matched_idx]
            if len(row.cells) >= 2:
                row.cells[1].text = value
        else:
            new_row = table.add_row()
            new_row.cells[0].text = field
            if len(new_row.cells) >= 2:
                new_row.cells[1].text = value


def _fill_two_column(table: Table, fields: list[str], data: dict):
    """
    Fill a two_column table where row 1 has field names as column headers
    and row 2 has their values. (e.g., Scope: "In Scope" | "Out of Scope")
    """
    if len(table.rows) < 2:
        return

    header_row = table.rows[1] if len(table.rows) > 1 else None
    if not header_row:
        return

    if len(table.rows) > 2:
        data_row = table.rows[2]
    else:
        data_row = table.add_row()

    for col_idx, field in enumerate(fields):
        value = str(data.get(field, "[TBD]"))
        if col_idx < len(data_row.cells):
            data_row.cells[col_idx].text = value


def _fill_text_block(table: Table, text: str):
    """
    Fill a text_block table where row 1 is a single merged cell
    containing free-form paragraph text.
    """
    if len(table.rows) < 2:
        table.add_row()

    row = table.rows[1]
    row.cells[0].text = str(text) if text else "[TBD]"


def fill_template(template_type: str, data: dict) -> bytes:
    """
    Open the .docx template, fill each section table with data from
    the LLM-generated JSON, and return the result as bytes.
    """
    schema = get_schema(template_type)
    template_path = get_template_path(template_type)
    doc = Document(str(template_path))

    for section in schema["sections"]:
        table = _find_table(doc, section["header_match"])
        if table is None:
            continue

        sec_id = section["id"]
        sec_data = data.get(sec_id)
        if sec_data is None:
            continue

        sec_type = section["type"]

        if sec_type == "columnar":
            rows_data = sec_data if isinstance(sec_data, list) else []
            _fill_columnar(table, section["columns"], rows_data)

        elif sec_type == "key_value":
            kv_data = sec_data if isinstance(sec_data, dict) else {}
            _fill_key_value(table, section.get("fields", []), kv_data)

        elif sec_type == "two_column":
            tc_data = sec_data if isinstance(sec_data, dict) else {}
            _fill_two_column(table, section.get("fields", []), tc_data)

        elif sec_type == "text_block":
            text = sec_data if isinstance(sec_data, str) else str(sec_data)
            _fill_text_block(table, text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def fill_and_save_template(
    project_id: int, template_type: str, data: dict
) -> str:
    """
    Fill template with data, save to project export folder, return filename.
    Uses same storage path as export_service.
    """
    from app.services.export_service import _get_documents_path

    file_bytes = fill_template(template_type, data)

    base = _get_documents_path()
    project_dir = base / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)

    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique = uuid.uuid4().hex[:8]
    filename = f"{template_type}_{stamp}_{unique}.docx"
    file_path = project_dir / filename
    file_path.write_bytes(file_bytes)

    return filename
