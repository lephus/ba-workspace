"""
Convert Markdown to DOCX/XLSX/MD and save to project folder.
Used for conversation export: save file to data/documents/<project_id>, return filename for download URL.
"""
import io
import re
import uuid
from datetime import datetime
from pathlib import Path

from flask import current_app

from app.services.config_loader import get_config

EXPORT_EXT = {"docx": "docx", "xlsx": "xlsx", "md": "md"}
EXPORT_MIME = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "md": "text/markdown",
}


def _get_documents_path() -> Path:
    config = get_config()
    path = config.get("documents_path")
    if not path:
        path = Path(current_app.config["PROJECT_ROOT"]) / "data" / "documents"
    return Path(path)


def _md_to_docx_bytes(markdown: str) -> bytes:
    """Convert Markdown to DOCX using python-docx (reliable, no external binary)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(stripped)
        elif stripped.startswith("|") and "|" in stripped[1:]:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells and not all(re.match(r"^[-:\s]+$", c) for c in cells):
                table = doc.add_table(rows=1, cols=len(cells))
                for j, cell in enumerate(cells):
                    table.rows[0].cells[j].text = cell
                i += 1
                while i < len(lines) and lines[i].strip().startswith("|"):
                    row_cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                    if row_cells and not all(re.match(r"^[-:\s]+$", c) for c in row_cells):
                        row = table.add_row()
                        for j, c in enumerate(row_cells):
                            if j < len(row.cells):
                                row.cells[j].text = c
                    i += 1
                i -= 1
        else:
            doc.add_paragraph(stripped)
        i += 1
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _md_to_xlsx_bytes(markdown: str) -> bytes:
    """Convert Markdown to XLSX: tables as sheets, rest as Content column (openpyxl)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()
    lines = markdown.split("\n")
    table_rows = []
    in_table = False
    for line in lines:
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            if cells and not all(re.match(r"^[-:\s]+$", c) for c in cells):
                table_rows.append(cells)
                in_table = True
            continue
        if in_table and table_rows:
            ws = wb.active if len(wb.worksheets) == 1 and wb.active.max_row == 0 else wb.create_sheet("Table")
            for r, row in enumerate(table_rows, 1):
                for c, val in enumerate(row, 1):
                    cell = ws.cell(row=r, column=c, value=val)
                    if r == 1:
                        cell.font = Font(bold=True)
            table_rows = []
        in_table = False
    if table_rows:
        ws = wb.active if wb.active.max_row == 0 else wb.create_sheet("Table")
        for r, row in enumerate(table_rows, 1):
            for c, val in enumerate(row, 1):
                ws.cell(row=r, column=c, value=val)
    if wb.active.max_row == 0:
        wb.active.title = "Content"
        for r, line in enumerate(markdown.split("\n"), 1):
            wb.active.cell(row=r, column=1, value=line)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def convert_markdown_to_bytes(markdown: str, fmt: str) -> bytes:
    """Convert Markdown to file bytes. fmt: docx, xlsx, md."""
    fmt = fmt.strip().lower()
    if fmt not in EXPORT_EXT:
        raise ValueError(f"Unsupported format: {fmt}. Allowed: {list(EXPORT_EXT.keys())}")
    if fmt == "md":
        return markdown.encode("utf-8")
    if fmt == "docx":
        return _md_to_docx_bytes(markdown)
    return _md_to_xlsx_bytes(markdown)


def save_export_to_project(project_id: int, markdown: str, fmt: str) -> str:
    """
    Convert Markdown to file, save under data/documents/<project_id>/export_<date>_<id>.<ext>.
    Returns filename (not full path) for use in download URL.
    File is removed when project is deleted (folder cleanup).
    """
    fmt = fmt.strip().lower()
    if fmt not in EXPORT_EXT:
        raise ValueError(f"Unsupported format: {fmt}")
    ext = EXPORT_EXT[fmt]
    base = _get_documents_path()
    project_dir = base / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique = uuid.uuid4().hex[:8]
    filename = f"export_{stamp}_{unique}.{ext}"
    file_path = project_dir / filename
    data = convert_markdown_to_bytes(markdown, fmt)
    file_path.write_bytes(data)
    return filename


def is_export_filename_safe(filename: str) -> bool:
    """Allow only export_*.docx|.xlsx|.md to prevent path traversal and arbitrary file access."""
    if not filename or ".." in filename or "/" in filename or "\\" in filename:
        return False
    if not filename.startswith("export_"):
        return False
    return any(filename.endswith("." + e) for e in EXPORT_EXT.values())
